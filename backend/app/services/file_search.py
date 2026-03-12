import re
import logging
from datetime import datetime
from typing import Optional
from pathlib import Path

from ..models.file import FileInfo, FileType, FileReadResponse
from ..models.search import SearchResult
from .file_scanner import FileScanner

logger = logging.getLogger(__name__)

# Extensions that can be summarized (text content can be extracted)
SUMMARIZABLE_EXTENSIONS = [
    # Text files (already supported)
    ".txt", ".md", ".json", ".xml", ".csv", ".log",
    ".py", ".js", ".ts", ".html", ".css", ".yaml", ".yml",
    ".sh", ".bash", ".zsh", ".conf", ".cfg", ".ini",
    # Document files (new support)
    ".pdf", ".docx", ".doc", ".xlsx", ".xls",
]

# Unsupported file types for summarization
UNSUPPORTED_EXTENSIONS = [
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".svg", ".webp", ".ico", ".tiff",  # Images
    ".mp4", ".avi", ".mkv", ".mov", ".wmv", ".flv", ".webm", ".m4v",  # Videos
    ".mp3", ".wav", ".ogg", ".flac", ".aac", ".m4a", ".wma",  # Audio
    ".zip", ".rar", ".7z", ".tar", ".gz", ".bz2", ".xz",  # Archives
    ".exe", ".dll", ".so", ".dylib", ".bin",  # Binaries
]


class FileSearchService:
    """Service for searching and reading files"""

    def __init__(self, scanner: FileScanner, device_name: str):
        self.scanner = scanner
        self.device_name = device_name

    def search(
        self,
        query: str,
        file_types: Optional[list[str]] = None,
        max_results: int = 50
    ) -> list[SearchResult]:
        """
        Search files by query.

        Supports:
        - Filename matching
        - Path matching
        - File extension filtering
        - Basic content search for indexed text
        """
        results = []

        # Convert file type strings to FileType enums
        type_filter = None
        if file_types:
            type_filter = []
            for ft in file_types:
                ft_lower = ft.lower().lstrip(".")
                # Map common extensions to file types
                extension_to_type = {
                    "pdf": FileType.DOCUMENT,
                    "doc": FileType.DOCUMENT,
                    "docx": FileType.DOCUMENT,
                    "txt": FileType.DOCUMENT,
                    "jpg": FileType.IMAGE,
                    "jpeg": FileType.IMAGE,
                    "png": FileType.IMAGE,
                    "mp4": FileType.VIDEO,
                    "mp3": FileType.AUDIO,
                }
                if ft_lower in extension_to_type:
                    type_filter.append(extension_to_type[ft_lower])

        # Search the index
        matched_files = self.scanner.search(query, type_filter)

        # Score and rank results
        query_lower = query.lower()
        query_words = set(re.findall(r"\w+", query_lower))

        for file_info in matched_files:
            score = self._calculate_relevance_score(file_info, query_lower, query_words)
            match_reason = self._get_match_reason(file_info, query_lower)

            result = SearchResult(
                file_info=file_info,
                device_name=self.device_name,
                device_id=self.scanner.device_id,
                relevance_score=score,
                match_reason=match_reason,
            )
            results.append(result)

        # Sort by relevance score
        results.sort(key=lambda x: x.relevance_score, reverse=True)

        return results[:max_results]

    def _calculate_relevance_score(
        self,
        file_info: FileInfo,
        query: str,
        query_words: set
    ) -> float:
        """Calculate relevance score for a file (0.0 to 1.0)"""
        score = 0.0
        name_lower = file_info.name.lower()

        # Exact filename match (highest score)
        if query == name_lower:
            return 1.0

        # Filename starts with query
        if name_lower.startswith(query):
            score += 0.8
        # Filename contains query
        elif query in name_lower:
            score += 0.6

        # Word matches in filename
        name_words = set(re.findall(r"\w+", name_lower))
        word_overlap = len(query_words & name_words) / max(len(query_words), 1)
        score += word_overlap * 0.3

        # Extension match bonus
        if file_info.extension.lstrip(".") in query:
            score += 0.1

        # Recent file bonus
        if file_info.modified_time:
            days_since_modified = (datetime.now() - file_info.modified_time).days
            if days_since_modified < 7:
                score += 0.1
            elif days_since_modified < 30:
                score += 0.05

        return min(score, 1.0)

    def _get_match_reason(self, file_info: FileInfo, query: str) -> str:
        """Explain why this file matched the query"""
        reasons = []

        if query in file_info.name.lower():
            reasons.append(f"Filename contains '{query}'")

        if query in file_info.path.lower() and query not in file_info.name.lower():
            reasons.append(f"Path contains '{query}'")

        if file_info.preview_text and query in file_info.preview_text.lower():
            reasons.append("Content contains search terms")

        return "; ".join(reasons) if reasons else "Filename or path matches"

    async def read_file(
        self,
        file_path: str,
        max_chars: int = 15000
    ) -> FileReadResponse:
        """Read the contents of a file for summarization"""
        file_info = self.scanner.get_file(file_path)

        if not file_info:
            return FileReadResponse(
                file_info=None,
                content=None,
                error="File not found in index",
            )

        ext = file_info.extension.lower()
        if not ext.startswith("."):
            ext = "." + ext

        # Check if file type is unsupported
        if ext in UNSUPPORTED_EXTENSIONS:
            file_type_name = self._get_unsupported_type_name(ext)
            return FileReadResponse(
                file_info=file_info,
                content=None,
                error=f"Cannot summarize {file_type_name} files ({ext}). This file type is not supported for summarization.",
            )

        # Check if file is summarizable
        if ext not in SUMMARIZABLE_EXTENSIONS:
            return FileReadResponse(
                file_info=file_info,
                content=None,
                error=f"Cannot read {ext} files. This file type is not supported for summarization.",
            )

        try:
            path = Path(file_path)
            if not path.exists():
                return FileReadResponse(
                    file_info=file_info,
                    content=None,
                    error="File does not exist on disk",
                )

            # Handle different file types
            content = None

            # PDF files
            if ext == ".pdf":
                content = await self._read_pdf(path)

            # DOCX files
            elif ext in [".docx", ".doc"]:
                content = await self._read_docx(path)

            # XLSX/XLS files
            elif ext in [".xlsx", ".xls"]:
                content = await self._read_xlsx(path)

            # Plain text files
            else:
                content = await self._read_text_file(path)

            if content is None:
                return FileReadResponse(
                    file_info=file_info,
                    content=None,
                    error="Failed to extract content from file",
                )

            truncated = len(content) > max_chars
            if truncated:
                content = content[:max_chars] + "\n\n... [content truncated for summarization]"

            return FileReadResponse(
                file_info=file_info,
                content=content,
                encoding="utf-8",
                truncated=truncated,
            )

        except PermissionError:
            return FileReadResponse(
                file_info=file_info,
                content=None,
                error="Permission denied",
            )
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {e}")
            return FileReadResponse(
                file_info=file_info,
                content=None,
                error=f"Error reading file: {str(e)}",
            )

    def _get_unsupported_type_name(self, ext: str) -> str:
        """Get human-readable name for unsupported file type"""
        if ext in [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".svg", ".webp", ".ico", ".tiff"]:
            return "image"
        elif ext in [".mp4", ".avi", ".mkv", ".mov", ".wmv", ".flv", ".webm", ".m4v"]:
            return "video"
        elif ext in [".mp3", ".wav", ".ogg", ".flac", ".aac", ".m4a", ".wma"]:
            return "audio"
        elif ext in [".zip", ".rar", ".7z", ".tar", ".gz", ".bz2", ".xz"]:
            return "archive"
        elif ext in [".exe", ".dll", ".so", ".dylib", ".bin"]:
            return "binary"
        return "unsupported"

    async def _read_text_file(self, path: Path) -> Optional[str]:
        """Read plain text file content"""
        import aiofiles
        async with aiofiles.open(path, "r", encoding="utf-8", errors="replace") as f:
            return await f.read()

    async def _read_pdf(self, path: Path) -> Optional[str]:
        """Extract text content from PDF file"""
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))

            text_parts = []
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")

            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning("pypdf not installed, cannot read PDF")
            return None
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            return None

    async def _read_docx(self, path: Path) -> Optional[str]:
        """Extract text content from DOCX file"""
        try:
            from docx import Document
            doc = Document(str(path))

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n".join(text_parts)
        except ImportError:
            logger.warning("python-docx not installed, cannot read DOCX")
            return None
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            return None

    async def _read_xlsx(self, path: Path) -> Optional[str]:
        """Extract text content from XLSX/XLS file"""
        try:
            from openpyxl import load_workbook
            wb = load_workbook(str(path), read_only=True, data_only=True)

            text_parts = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"=== Sheet: {sheet_name} ===")

                for row in sheet.iter_rows(max_row=100):  # Limit rows
                    row_values = []
                    for cell in row:
                        if cell.value is not None:
                            row_values.append(str(cell.value))
                    if row_values:
                        text_parts.append(" | ".join(row_values))

            return "\n".join(text_parts)
        except ImportError:
            logger.warning("openpyxl not installed, cannot read XLSX")
            return None
        except Exception as e:
            logger.error(f"Error reading XLSX: {e}")
            return None

    def get_file_path(self, file_path: str) -> Optional[Path]:
        """Get the Path object for a file, with validation"""
        file_info = self.scanner.get_file(file_path)
        if not file_info:
            return None

        path = Path(file_path)
        if not path.exists():
            return None

        return path
